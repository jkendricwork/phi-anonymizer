"""Document parsing service for Word and PDF files."""

import os
from pathlib import Path
from typing import Tuple

from app.core.config import settings


async def parse_word(file_path: str) -> str:
    """
    Extract text from Word document (.docx).

    Args:
        file_path: Path to the .docx file

    Returns:
        Extracted text with basic formatting preserved
    """
    from docx import Document

    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    return "\n".join(text_parts)


def detect_needs_ocr(pdf_path: str) -> bool:
    """
    Check if PDF needs OCR (is scanned/image-based).

    Args:
        pdf_path: Path to the PDF file

    Returns:
        True if OCR is needed, False otherwise
    """
    import fitz  # PyMuPDF

    doc = fitz.open(pdf_path)

    # Check first few pages for text content
    pages_to_check = min(3, len(doc))
    total_text_length = 0

    for page_num in range(pages_to_check):
        page = doc[page_num]
        text = page.get_text()
        total_text_length += len(text.strip())

    doc.close()

    # If average text per page is very low, likely needs OCR
    avg_text_per_page = total_text_length / pages_to_check if pages_to_check > 0 else 0
    return avg_text_per_page < 100  # Threshold for text content


async def parse_pdf_text(file_path: str) -> str:
    """
    Extract text from PDF using PyMuPDF.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    text_parts = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()

        if text.strip():
            text_parts.append(f"--- Page {page_num + 1} ---")
            text_parts.append(text.strip())

    doc.close()

    return "\n\n".join(text_parts)


async def parse_pdf_ocr(file_path: str) -> str:
    """
    Extract text from PDF using OCR (for scanned documents).

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text from OCR
    """
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image

    # Set Tesseract path if configured
    if settings.TESSERACT_PATH and os.path.exists(settings.TESSERACT_PATH):
        pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH

    # Convert PDF to images
    images = convert_from_path(file_path, dpi=300)

    text_parts = []

    for i, image in enumerate(images):
        # Perform OCR on each page
        text = pytesseract.image_to_string(image, lang=settings.TESSERACT_LANG)

        if text.strip():
            text_parts.append(f"--- Page {i + 1} ---")
            text_parts.append(text.strip())

    return "\n\n".join(text_parts)


async def parse_pdf(file_path: str, force_ocr: bool = False) -> Tuple[str, bool]:
    """
    Extract text from PDF, automatically detecting if OCR is needed.

    Args:
        file_path: Path to the PDF file
        force_ocr: Force OCR even if text is detected

    Returns:
        Tuple of (extracted_text, used_ocr)
    """
    # Check if OCR is needed
    needs_ocr = force_ocr or detect_needs_ocr(file_path)

    if needs_ocr:
        text = await parse_pdf_ocr(file_path)
        return text, True
    else:
        text = await parse_pdf_text(file_path)
        return text, False


async def parse_document(file_path: str, file_extension: str) -> Tuple[str, bool]:
    """
    Parse document based on file type.

    Args:
        file_path: Path to the document file
        file_extension: File extension (.docx, .pdf)

    Returns:
        Tuple of (extracted_text, used_ocr)

    Raises:
        ValueError: If file type is not supported
    """
    file_extension = file_extension.lower()

    if file_extension == ".docx":
        text = await parse_word(file_path)
        return text, False

    elif file_extension == ".pdf":
        return await parse_pdf(file_path)

    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
